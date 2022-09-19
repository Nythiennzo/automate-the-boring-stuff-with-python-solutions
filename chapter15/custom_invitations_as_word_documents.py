import docx

guests_list_file_path = r'C:\repos\python\automate_the_boring_stuff\chapter15\guests.txt'
template_docx_path = r'C:\repos\python\automate_the_boring_stuff\chapter15\Invitation.docx'
target_docx_path = r'C:\repos\python\automate_the_boring_stuff\chapter15\InvitationWithGuests.docx'

guests_list_file = open(guests_list_file_path)
guests = [guest.rstrip('\n') for guest in guests_list_file.readlines()]
guests_list_file.close()

doc = docx.Document(template_docx_path)

template_lenth = len(doc.paragraphs)

for guest in guests:
    doc.add_page_break()
    for paragraph_index in range(template_lenth):
        source_paragraph = doc.paragraphs[paragraph_index]
        writing_paragraph = doc.add_paragraph('', source_paragraph.style)

        for source_run in source_paragraph.runs:
            writing_text = source_run.text.replace('<guest>', guest)
            writing_run = writing_paragraph.add_run(writing_text)

            writing_run.font.name = source_run.font.name
            writing_run.font.size = source_run.font.size

        writing_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment

doc.save(target_docx_path)

